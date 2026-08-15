from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[1]
BUILDER = PROJECT_ROOT / "tools" / "build_documents.py"


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class DocumentDeliveryTests(unittest.TestCase):
    def test_builder_creates_complete_manual_and_one_page_resume(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output_dir = Path(directory)
            completed = subprocess.run(
                [sys.executable, str(BUILDER), "--output-dir", str(output_dir)],
                cwd=PROJECT_ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(completed.returncode, 0, completed.stderr)

            manual = output_dir / "蔡健明_漫剧投放数据项目_实操手册.docx"
            resume_docx = output_dir / "蔡健明_字节数据岗位_校招简历.docx"
            resume_pdf = output_dir / "蔡健明_字节数据岗位_校招简历.pdf"
            self.assertTrue(manual.exists())
            self.assertTrue(resume_docx.exists())
            self.assertTrue(resume_pdf.exists())

            manual_text = document_text(manual)
            for required in (
                "Excel 零基础操作",
                "MySQL 8 导入与 SQL 执行",
                "Python 管道与 Notebook",
                "Power BI 三页看板",
                "GitHub Pages 部署",
                "完整代码附录",
                "def run_pipeline",
                "eligible_ordered AS",
            ):
                self.assertIn(required, manual_text)

            resume_text = document_text(resume_docx)
            for required in (
                "蔡健明",
                "19012716053",
                "1764798435@qq.com",
                "杭州",
                "浙江财经大学",
                "经济统计学",
                "2023.09–2027.06",
                "三次历史周数据复盘",
                "个人分析项目",
                "浙江省城乡居民收入分析",
                "10+ 场",
            ):
                self.assertIn(required, resume_text)

            for forbidden in (
                "[请填写",
                "请按实习证明核对",
                "正式上线",
                "ROI 提升",
                "节省工时",
            ):
                self.assertNotIn(forbidden, resume_text)

            reader = PdfReader(resume_pdf)
            self.assertEqual(len(reader.pages), 1)
            extracted = reader.pages[0].extract_text() or ""
            self.assertIn("蔡健明", extracted)
            self.assertIn("19012716053", extracted)
            self.assertNotIn("[请填写", extracted)


if __name__ == "__main__":
    unittest.main()
