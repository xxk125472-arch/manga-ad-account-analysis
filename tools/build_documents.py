#!/usr/bin/env python3
"""Build the beginner operating manual and the one-page targeted resume.

The manual copies code directly from this repository so that its appendices
cannot silently drift away from the implementation being delivered.
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "deliverables"
FONT_DIR = Path("/tmp/codex-fontpkg/node_modules/@fontpkg/noto-sans-cjk-sc").resolve()
BODY_FONT = "Noto Sans CJK SC"
CODE_FONT = "DejaVu Sans Mono"

NAVY = "1B2A4A"
BLUE = "3867D6"
TEAL = "148A8B"
CORAL = "E86A4A"
GOLD = "C98A22"
INK = "17213A"
MUTED = "667085"
LINE = "D8DEE8"
BLUE_GRAY = "E8EEF5"
LIGHT = "F4F6F9"
WHITE = "FFFFFF"
SOFT_GOLD = "FFF6DF"


MANUAL_NAME = "蔡健明_漫剧投放数据项目_实操手册.docx"
RESUME_DOCX_NAME = "蔡健明_字节数据岗位_校招简历.docx"
RESUME_PDF_NAME = "蔡健明_字节数据岗位_校招简历.pdf"


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    size: float = 11,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = name
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color: str = LINE, size: str = "4") -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_in: list[float], *, indent_dxa: int = 120) -> None:
    widths_dxa = [round(value * 1440) for value in widths_in]
    total = sum(widths_dxa)
    table.autofit = False
    properties = table._tbl.tblPr

    width_node = properties.first_child_found_in("w:tblW")
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        properties.append(width_node)
    width_node.set(qn("w:w"), str(total))
    width_node.set(qn("w:type"), "dxa")

    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8, color=MUTED)


def add_paragraph_border(paragraph, *, color: str = CORAL, size: str = "14") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def configure_manual_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specifications = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, MUTED, 0, 8),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in specifications.items():
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = style_name.startswith("Heading")

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_manual_section(section) -> None:
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run("漫剧投放账号效能诊断 · 实操手册"), size=8, color=MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run("蔡健明  |  "), size=8, color=MUTED)
    add_page_field(paragraph)


def add_body(
    document: Document,
    text: str,
    *,
    bold_prefix: str | None = None,
    size: float = 11,
    color: str = INK,
    italic: bool = False,
    after: float = 6,
) -> object:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(paragraph.add_run(bold_prefix), size=size, color=color, bold=True)
        set_run_font(paragraph.add_run(text[len(bold_prefix) :]), size=size, color=color, italic=italic)
    else:
        set_run_font(paragraph.add_run(text), size=size, color=color, italic=italic)
    return paragraph


def add_bullet(document: Document, text: str, *, size: float = 10.6) -> object:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text), size=size)
    return paragraph


def add_step(document: Document, text: str, *, size: float = 10.6) -> object:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text), size=size)
    return paragraph


def add_code_block(document: Document, code: str, *, language: str = "text") -> None:
    label = document.add_paragraph()
    label.paragraph_format.space_before = Pt(3)
    label.paragraph_format.space_after = Pt(2)
    set_run_font(label.add_run(language.upper()), size=7.2, color=MUTED, bold=True)
    for line in code.rstrip("\n").splitlines() or [""]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.right_indent = Inches(0.08)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        properties = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F3F5F7")
        properties.append(shading)
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, name=CODE_FONT, size=6.5, color="243247")
        # Keep exact source lines copyable while preventing long Python/SQL/DAX
        # statements from crossing the printable A4 text area.
        scale = OxmlElement("w:w")
        scale.set(qn("w:val"), "75")
        run._element.get_or_add_rPr().append(scale)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_callout(document: Document, title: str, body: str, *, tone: str = "blue") -> None:
    palette = {
        "blue": (BLUE, "EFF3FF"),
        "teal": (TEAL, "EAF7F5"),
        "gold": (GOLD, SOFT_GOLD),
        "coral": (CORAL, "FDECE7"),
    }
    accent, fill = palette[tone]
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [1.25, 5.25])
    set_table_borders(table, color=accent, size="6")
    for cell in table.rows[0].cells:
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    set_run_font(left.add_run(title), size=10.2, color=accent, bold=True)
    right = table.cell(0, 1).paragraphs[0]
    right.paragraph_format.space_after = Pt(0)
    right.paragraph_format.line_spacing = 1.25
    set_run_font(right.add_run(body), size=9.2)
    add_body(document, "", after=2)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_in: list[float],
    *,
    font_size: float = 8.5,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths_in)
    set_table_borders(table)
    header = table.rows[0]
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, BLUE_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        set_run_font(paragraph.add_run(value), size=font_size, color=NAVY, bold=True)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header._tr.get_or_add_trPr().append(repeat)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cell, "FAFBFC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            set_run_font(paragraph.add_run(str(value)), size=font_size)
    add_body(document, "", after=2)


def add_heading(document: Document, text: str, level: int = 1) -> object:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    set_run_font(
        paragraph.add_run(text),
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: "1F4D78"}[level],
        bold=True,
    )
    return paragraph


def add_manual_cover(document: Document) -> None:
    for _ in range(3):
        add_body(document, "", after=14)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    set_run_font(kicker.add_run("PORTFOLIO OPERATING GUIDE · V1.0"), size=9, color=CORAL, bold=True)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("漫剧投放账号效能诊断\n与预算分层管理"), size=30, color=NAVY, bold=True)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(subtitle.add_run("零基础复现手册｜Excel + MySQL + Python + Power BI + GitHub Pages"), size=13.5, color=MUTED)

    add_callout(
        document,
        "一句话目标",
        "从深度脱敏的 2026 年 7 月账号月度数据出发，亲手复现质量审计、加权口径、Pareto、五类行动池、三端看板与 GitHub 作品集。",
        tone="teal",
    )

    add_table(
        document,
        ["作者", "数据窗口", "样本", "交付形态"],
        [["蔡健明", "2026-07", "551 账号｜8 品牌", "GitHub-ready 项目 + 手册 + 简历"]],
        [1.1, 1.2, 1.6, 2.6],
        font_size=8.8,
    )
    add_body(
        document,
        "诚信边界：这是基于真实实习业务场景主动发起的个人分析项目，不是公司正式上线系统。金额经过匿名替换与同比例缩放；单月账号粒度不能证明趋势、归因、LTV、利润或策略收益。",
        size=9.5,
        color=MUTED,
        italic=True,
        after=10,
    )
    document.add_page_break()


def add_image(document: Document, relative_path: str, *, width: float = 6.35, caption: str = "") -> None:
    path = PROJECT_ROOT / relative_path
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(6)
        set_run_font(caption_paragraph.add_run(caption), size=8, color=MUTED)


def build_manual(output_path: Path) -> None:
    document = Document()
    configure_manual_styles(document)
    for section in document.sections:
        configure_manual_section(section)
    properties = document.core_properties
    properties.title = "漫剧投放账号效能诊断与预算分层管理实操手册"
    properties.subject = "GitHub 数据分析作品集零基础复现"
    properties.author = "蔡健明"
    properties.last_modified_by = "蔡健明"
    properties.keywords = "数据分析, SQL, Python, Power BI, GitHub Pages"

    add_manual_cover(document)

    add_heading(document, "如何使用这份手册", 1)
    add_body(document, "这份手册按“先复现、再理解、最后表达”的顺序编排。第一次不要跳步骤；每一阶段都留下可截图、可对账、可在面试中解释的证据。")
    for item in (
        "全程只使用仓库 data/processed 中的公开脱敏数据；不要把原始未脱敏表放进仓库。",
        "看到命令时，先确认终端当前位于项目根目录；路径错误是初学者最常见的问题。",
        "所有品牌与全局比率都先求和分子分母再相除，不平均账号级 CTR、CPC 或 ROI。",
        "代码必须实际运行；看到报错先复制完整报错，再到本手册“故障排查”按症状处理。",
        "面试只说完成的交付与观察结果，不声称上线、增效、利润或因果。",
    ):
        add_bullet(document, item)

    add_heading(document, "30 分钟最短复现路线", 2)
    for step in (
        "解压项目并在根目录打开终端。",
        "创建 Python 虚拟环境，安装依赖并运行 scripts/run_pipeline.py。",
        "执行全部单元测试与两个 Notebook，确认 551 行、8 品牌、283 个有效账号。",
        "进入 web/，执行 npm ci、npm test、npm run dev，检查三个页签与筛选联动。",
        "阅读 docs/analysis_report.pdf 与 README.md，练习 30 秒和 3 分钟讲解。",
    ):
        add_step(document, step)
    add_code_block(
        document,
        """python -m venv .venv
# Windows PowerShell
.venv\\Scripts\\Activate.ps1
python -m pip install --upgrade pip
pip install -e \".[dev]\"
python scripts/run_pipeline.py
python scripts/build_notebooks.py
python scripts/execute_notebooks.py
PYTHONPATH=src python -m unittest discover -s tests -v

cd web
npm ci
npm test
npm run dev""",
        language="shell",
    )

    document.add_page_break()
    add_heading(document, "目录", 1)
    for item in (
        "1. 项目定位、数据契约与真实边界",
        "2. 软件安装与环境验收",
        "3. Excel 零基础操作",
        "4. MySQL 8 导入与 SQL 执行",
        "5. Python 管道与 Notebook",
        "6. Power BI 三页看板",
        "7. React 看板、本地 QA 与 GitHub Pages",
        "8. 结果解读、行动框架与面试表达",
        "9. Git/GitHub 上传与更新",
        "10. 故障排查与最终验收",
        "附录 A. 完整代码附录",
    ):
        add_bullet(document, item)
    document.add_page_break()

    add_heading(document, "1. 项目定位、数据契约与真实边界", 1)
    add_heading(document, "1.1 项目要回答什么", 2)
    add_body(document, "业务问题不是“哪个单一指标万能”，而是：预算集中在哪里、规模与回收是否匹配、哪些账号应优先扩量/优化/清理、哪些账号因样本太小只能继续观察。")
    add_table(
        document,
        ["层级", "问题", "交付"],
        [
            ["数据质量", "字段、主键、缺失、零值与口径是否可信", "质量检查表 + 播放成本审计"],
            ["组合经营", "预算是否集中、头部承担多少消耗", "Pareto 与头部贡献"],
            ["品牌对标", "规模与加权 ROI 是否匹配", "品牌汇总与加权指标"],
            ["账号动作", "先处理谁、为什么、怎么验证", "五类互斥行动池"],
        ],
        [1.15, 2.65, 2.7],
    )

    add_heading(document, "1.2 已验证的数据契约", 2)
    add_table(
        document,
        ["检查项", "结果", "解释"],
        [
            ["时间", "2026-07", "只有一个自然月，不做趋势"],
            ["粒度", "账号 × 月", "每个账号一行"],
            ["规模", "551 行 / 551 唯一账号", "无重复账号月键"],
            ["品牌", "8 个匿名品牌", "不是 9 个；以源数据实测为准"],
            ["字段", "22 个源字段", "客户字段不进入公开网页 JSON"],
            ["完整性", "0 缺失 / 0 负值", "零分母单独处理，不等于缺失"],
        ],
        [1.1, 2.0, 3.4],
    )
    add_callout(
        document,
        "关键审计",
        "源字段“播放成本”在有消耗账号中实际匹配 播放量 ÷ 消耗。项目保留它做审计，但所有成本分析统一重算为 消耗 ÷ 播放量。",
        tone="coral",
    )

    add_heading(document, "1.3 Headline 对账基线", 2)
    add_table(
        document,
        ["指标", "基线", "口径"],
        [
            ["总消耗", "207,057.29", "SUM(platform_spend)"],
            ["24H 混合变现", "213,266.26", "SUM(mixed_revenue_24h)"],
            ["加权 24H 混合 ROI", "1.03", "总混合变现 ÷ 总消耗"],
            ["有效账号", "283", "消耗 ≥ 50"],
            ["有效样本消耗覆盖率", "99.43%", "有效账号消耗 ÷ 总消耗"],
            ["高消耗边界", "735.80", "有效账号消耗 P75"],
            ["Top 10%", "55 个账号 / 64.74% 消耗", "floor(551×10%)"],
        ],
        [2.0, 1.9, 2.6],
    )

    add_heading(document, "1.4 真实经历怎么衔接", 2)
    add_body(document, "公司内实际工作以视频素材上传、B 站推送、广告系统品牌/产品基础配置与计划搭建为主。你在日常工作外主动做过三次历史周数据复盘并向带教展示；当前完整仓库是在脱敏数据上继续完成的个人分析项目。")
    add_callout(
        document,
        "面试说法",
        "“我在执行型工作中发现账号多、口径杂、复盘难直接对应动作，因此在工作外用深度脱敏数据把问题继续做成可复现作品集。项目没有在公司正式上线，也没有验证 ROI 提升。”",
        tone="gold",
    )

    add_heading(document, "1.5 项目目录", 2)
    add_code_block(
        document,
        """manga-ad-account-analysis/
├─ README.md                  # 面试官先读这里
├─ data/processed/            # 深度脱敏源表与派生 CSV
├─ sql/                       # MySQL 8 建表、质量、品牌、Pareto、分层
├─ src/manga_ad_analysis/     # 唯一业务逻辑来源
├─ scripts/                   # 管道与 Notebook 入口
├─ notebooks/                 # 已执行 Notebook
├─ outputs/                   # 管道验证回执
├─ images/                    # 静态图与看板截图
├─ docs/                      # 数据字典、报告、面试与部署说明
├─ powerbi/                   # DAX、M、主题、三页施工清单
├─ web/                       # React/Vite 静态交互看板
├─ tests/                     # 数据与交付契约测试
└─ .github/workflows/         # Pages 自动部署""",
        language="text",
    )

    add_heading(document, "2. 软件安装与环境验收", 1)
    add_heading(document, "2.1 建议版本", 2)
    add_table(
        document,
        ["软件", "建议版本", "只需掌握"],
        [
            ["Excel", "Microsoft 365 / 2021+", "表格、筛选、公式、CSV"],
            ["Python", "3.11 或 3.12", "venv、pip、pandas、Notebook"],
            ["Git", "2.40+", "init/add/commit/push"],
            ["Node.js", "20 LTS+", "npm ci/test/build/dev"],
            ["MySQL", "8.0+", "建表、导入、CTE、窗口函数"],
            ["Power BI Desktop", "当期稳定版（Windows）", "Power Query、DAX、三页看板"],
            ["VS Code", "当期稳定版", "终端、Python、SQL、Notebook"],
        ],
        [1.3, 1.6, 3.6],
    )
    add_heading(document, "2.2 Windows 安装顺序", 2)
    for step in (
        "安装 Git：从 git-scm.com 下载，保持默认选项；完成后重新打开 PowerShell。",
        "安装 Python：勾选 Add Python to PATH；安装后执行 python --version。",
        "安装 Node.js LTS：完成后执行 node --version 与 npm --version。",
        "安装 MySQL Community Server + Workbench：记录 root 密码，不写进仓库。",
        "安装 Power BI Desktop：优先 Microsoft Store 或微软官网下载；macOS 无原生 Desktop。",
        "安装 VS Code，并添加 Python、Jupyter、SQLTools（可选）扩展。",
    ):
        add_step(document, step)
    add_code_block(
        document,
        """git --version
python --version
python -m pip --version
node --version
npm --version
mysql --version""",
        language="powershell",
    )

    add_heading(document, "2.3 初始化 Python 环境", 2)
    for step in (
        "在项目根目录地址栏输入 powershell 并回车。",
        "创建 .venv；它把依赖隔离在项目内部。",
        "激活虚拟环境；终端行首出现 (.venv) 即成功。",
        "安装项目与开发依赖。不要把 .venv 上传 GitHub。",
    ):
        add_step(document, step)
    add_code_block(
        document,
        """python -m venv .venv
.venv\\Scripts\\Activate.ps1
python -m pip install --upgrade pip
pip install -e \".[dev]\"""",
        language="powershell",
    )
    add_callout(
        document,
        "脚本限制",
        "若 PowerShell 提示“禁止运行脚本”，只对当前窗口执行 Set-ExecutionPolicy -Scope Process Bypass，然后重新激活。不要修改公司电脑的全局安全策略。",
        tone="gold",
    )

    add_heading(document, "3. Excel 零基础操作", 1)
    add_heading(document, "3.1 打开并确认源文件", 2)
    for step in (
        "打开 data/processed/manga_ad_account_2026_07_anonymized.xlsx。若只拿到原附件，先另存到该英文路径。",
        "只分析“脱敏数据”工作表；“脱敏说明”和“字段说明”用于解释，不参与聚合。",
        "单击任意单元格后按 Ctrl+End，确认数据末行约为 552（第 1 行是表头）。",
        "选择 A1:V552，按 Ctrl+T，勾选“表包含标题”，把表名改为 源数据。",
        "视图 → 冻结窗格 → 冻结首行；数据 → 筛选，检查日期、品牌与零值。",
    ):
        add_step(document, step)

    add_heading(document, "3.2 数据类型与筛选", 2)
    add_table(
        document,
        ["字段组", "Excel 类型", "操作"],
        [
            ["日期", "日期", "设置 yyyy-mm-dd；筛选应只有 2026-07-01"],
            ["账号/客户/品牌", "文本", "不要转成数字；账号 ID 保留前导字符"],
            ["展示/点击/播放", "整数", "千位分隔，0 位小数"],
            ["消耗/变现/成本", "数值", "2–5 位小数，计算后再格式化"],
            ["CTR/ROI/播放率", "数值", "只作核对；聚合时从基础量重算"],
        ],
        [1.4, 1.3, 3.8],
    )

    add_heading(document, "3.3 五个必做公式", 2)
    add_body(document, "在“源数据”表右侧新增帮助列。若 Excel 使用中文分号分隔参数，请把公式中的逗号替换为分号。")
    formulas = [
        ("账号月重复计数", '=COUNTIFS(源数据[日期],[@日期],源数据[账号id],[@账号id])'),
        ("重算 CTR", '=IFERROR([@点击量]/[@展示量],\"\")'),
        ("重算播放率", '=IFERROR([@播放量]/[@展示量],\"\")'),
        ("重算播放成本", '=IFERROR([@[三连投放总消耗(毫分)]]/[@播放量],\"\")'),
        ("源播放成本反向审计", '=IF(ABS([@播放成本]-IFERROR([@播放量]/[@[三连投放总消耗(毫分)]],0))<0.000001,\"反向口径匹配\",\"需检查\")'),
    ]
    for name, formula in formulas:
        add_body(document, name, bold_prefix=name, after=2)
        add_code_block(document, formula, language="excel")

    add_heading(document, "3.4 加权指标：不能 AVERAGE", 2)
    add_body(document, "品牌或整体指标必须从基础量求和后再相除。假设源表名为“源数据”，使用：")
    add_code_block(
        document,
        """加权 CTR
=SUM(源数据[点击量])/SUM(源数据[展示量])

加权 CPC
=SUM(源数据[三连投放总消耗(毫分)])/SUM(源数据[点击量])

加权播放率
=SUM(源数据[播放量])/SUM(源数据[展示量])

重算播放成本
=SUM(源数据[三连投放总消耗(毫分)])/SUM(源数据[播放量])

24H 混合 ROI
=SUM(源数据[激活24小时内混合变现金额])/SUM(源数据[三连投放总消耗(毫分)])""",
        language="excel",
    )
    add_callout(
        document,
        "为什么",
        "平均账号 ROI 会让消耗 1 元与消耗 10,000 元的账号权重相同。总变现 ÷ 总消耗才回答“这一组合每 1 单位消耗带来多少变现”。",
        tone="teal",
    )

    add_heading(document, "3.5 门槛、P75 与行动标签", 2)
    add_code_block(
        document,
        """有效账号判断
=IF([@[三连投放总消耗(毫分)]]>=50,1,0)

有效账号消耗 P75（Microsoft 365）
=PERCENTILE.INC(FILTER(源数据[三连投放总消耗(毫分)],源数据[三连投放总消耗(毫分)]>=50),0.75)

行动标签（假设 P75 位于 $AA$2，重算 ROI 位于 [@重算ROI]）
=IF([@[三连投放总消耗(毫分)]]<50,\"数据不足/低量池\",
 IF([@[三连投放总消耗(毫分)]]>=$AA$2,
    IF([@重算ROI]>=1,\"核心扩量候选\",\"高消耗重点优化\"),
    IF([@重算ROI]>=1,\"小步扩量观察\",\"低效清理候选\")))""",
        language="excel",
    )
    add_body(document, "预期 P75 为 735.80492（显示为 735.80）。标签必须互斥、完备，五类合计 551。")

    add_heading(document, "3.6 导出 CSV", 2)
    for step in (
        "优先使用仓库已生成的 data/processed/manga_ad_account_2026_07.csv；它已改为英文列名并保留 UTF-8。",
        "若手工导出：文件 → 另存为 → CSV UTF-8（逗号分隔）。不要用旧版 ANSI CSV。",
        "Excel 只会保存当前工作表；看到提示时选择“使用 CSV 格式”。",
        "关闭后重新打开 CSV，确认中文、日期、账号 ID 与小数位未损坏。",
    ):
        add_step(document, step)

    add_heading(document, "4. MySQL 8 导入与 SQL 执行", 1)
    add_heading(document, "4.1 新建数据库", 2)
    add_code_block(
        document,
        """mysql -u root -p
CREATE DATABASE IF NOT EXISTS manga_ad_analysis
  CHARACTER SET utf8mb4
  COLLATE utf8mb4_0900_ai_ci;
USE manga_ad_analysis;""",
        language="sql",
    )
    add_heading(document, "4.2 按顺序执行脚本", 2)
    for step in (
        "在 MySQL Workbench 打开 sql/01_create_table.sql，执行并确认 fact_account_month 表创建。",
        "用 Table Data Import Wizard 导入 data/processed/manga_ad_account_2026_07.csv，目标表选择 fact_account_month。",
        "检查字段映射与日期格式；导入完成后 SELECT COUNT(*) 应为 551。",
        "依次执行 02_data_quality、03_brand_summary、04_account_pareto、05_account_segmentation。",
        "把查询结果与 data/processed 下的 CSV 对账；不要手工改 SQL 结果迎合结论。",
    ):
        add_step(document, step)
    add_code_block(
        document,
        """SELECT COUNT(*) AS rows,
       COUNT(DISTINCT account_id) AS accounts,
       COUNT(DISTINCT brand_name) AS brands,
       ROUND(SUM(platform_spend), 2) AS spend
FROM fact_account_month;

-- 预期：551 / 551 / 8 / 207057.29""",
        language="sql",
    )
    add_callout(
        document,
        "导入失败",
        "若 Workbench 把 UTF-8 BOM 当成列名的一部分，先使用导入向导而不是 LOAD DATA；仍失败时打开 CSV，另存为 UTF-8 无 BOM 后重试。仓库 SQL 的建表字段名必须与 CSV 英文表头一致。",
        tone="gold",
    )

    add_heading(document, "4.3 读懂五类 SQL", 2)
    add_table(
        document,
        ["脚本", "学习重点", "验收"],
        [
            ["01_create_table.sql", "数据类型、主键、索引", "表结构与 22 个字段映射"],
            ["02_data_quality.sql", "COUNT/CASE/NULLIF", "缺失、重复、零值与口径冲突"],
            ["03_brand_summary.sql", "GROUP BY + 加权比率", "8 个品牌；消耗合计不变"],
            ["04_account_pareto.sql", "窗口累计和/排名", "Top 55 = 64.74%"],
            ["05_account_segmentation.sql", "CTE、ROW_NUMBER、P75、CASE", "283 有效账号；五类合计 551"],
        ],
        [1.55, 2.45, 2.5],
    )

    add_heading(document, "5. Python 管道与 Notebook", 1)
    add_heading(document, "5.1 一键运行", 2)
    add_code_block(
        document,
        """# 项目根目录；Windows 已激活 .venv
python scripts/run_pipeline.py
python scripts/build_notebooks.py
python scripts/execute_notebooks.py
PYTHONPATH=src python -m unittest discover -s tests -v""",
        language="shell",
    )
    add_body(document, "成功后查看 outputs/pipeline_receipt.json；all_reconciliations_passed 必须为 true。若数字不符，先停止后续 Power BI 与网页工作，回到源文件与字段映射。")

    add_heading(document, "5.2 代码职责", 2)
    add_table(
        document,
        ["文件", "职责", "关键检查"],
        [
            ["io.py", "中文列名 → 英文列名、数据类型与源契约", "551×22、8 品牌、2026-07"],
            ["quality.py", "缺失/重复/负值/零分母/播放成本冲突", "不把 0 当缺失"],
            ["metrics.py", "安全除法、加权指标、品牌汇总、Pareto", "先求和再相除"],
            ["segmentation.py", "门槛、P75、五类互斥标签", "283 有效；P75=735.80"],
            ["pipeline.py", "输出 CSV/JSON/PNG/回执", "公开 JSON 无客户字段"],
            ["visuals.py", "四张静态图", "标题、单位、参考线与中文字体"],
        ],
        [1.35, 2.8, 2.35],
    )

    add_heading(document, "5.3 Notebook 怎么读", 2)
    for step in (
        "打开 notebooks/01_data_quality_audit.ipynb，从上到下运行，确认无红色异常堆栈。",
        "先读每个 Markdown 单元的“为什么”，再读代码；不要只看最后图表。",
        "打开 02_account_performance_diagnosis.ipynb，依次核对品牌、Pareto、行动池与相关性。",
        "相关系数绝对值小不代表数据无用；它说明不能用单一前链路指标解释回收。",
        "修改参数时先复制 Notebook；正式交付保留基线版本。",
    ):
        add_step(document, step)
    add_image(document, "images/account_pareto.png", caption="账号消耗 Pareto：Top 55 账号贡献 64.74% 消耗。")

    add_heading(document, "5.4 输出文件", 2)
    add_table(
        document,
        ["输出", "给谁用", "不要做什么"],
        [
            ["brand_summary.csv", "报告 / 品牌对标", "不要再平均其中的比例"],
            ["account_pareto.csv", "Pareto / 头部管理", "不要把集中度直接称为风险"],
            ["account_segments.csv", "Power BI / 行动池", "标签不是自动投放指令"],
            ["correlation_matrix.csv", "探索 / 后续问题", "不要写成因果或驱动"],
            ["dashboard.json", "React 静态看板", "不要加入客户字段"],
            ["pipeline_receipt.json", "跨端对账", "失败时不可继续发布"],
        ],
        [1.75, 2.05, 2.7],
    )

    add_heading(document, "6. Power BI 三页看板", 1)
    add_callout(
        document,
        "交付说明",
        "仓库提供完整施工包，不提供伪造 .pbix。请在 Windows Power BI Desktop 亲自按步骤建立模型并保存；面试时只说“设计模型、DAX 与施工包”，直到你确实生成并检查 PBIX。",
        tone="coral",
    )
    add_heading(document, "6.1 导入与星型模型", 2)
    for step in (
        "打开 Power BI Desktop → 获取数据 → 文本/CSV → 选择 data/processed/account_segments.csv。",
        "点击“转换数据”，确认 month_date 为日期；展示/点击/播放为整数；消耗、变现、成本、ROI 为小数。",
        "或按 powerbi/power_query.m 新建参数 pDataPath 与四个查询：账户投放、品牌、账号、日期。",
        "建立 品牌[品牌] 1:* 账户投放[brand_name]、账号[account_id] 1:* 账户投放[account_id]、日期[日期] 1:* 账户投放[month_date]。",
        "筛选方向均从维表到事实表；不要建多对多。把日期表标记为日期表，但不画单月趋势。",
    ):
        add_step(document, step)

    add_heading(document, "6.2 导入主题与 DAX", 2)
    for step in (
        "视图 → 主题 → 浏览主题 → 选择 powerbi/theme.json。",
        "主页 → 输入数据，新建一列占位并命名表为“指标”；随后隐藏占位列。",
        "逐段复制 powerbi/measures.dax。先建立基础总额，再建立比率与 QA 卡。",
        "设置格式：CTR/播放率为百分比 2 位；ROI 2 位；金额/成本 2 位；账号数为整数。",
        "无筛选时对账 551 / 8 / 207,057.29 / 1.03 / 283 / 99.43% / 735.80。",
    ):
        add_step(document, step)

    add_heading(document, "6.3 第 1 页：经营总览", 2)
    for item in (
        "顶部：品牌、行动标签、样本状态切片器；同步到三页。",
        "KPI：总消耗、24H 混合 ROI、加权 CTR、加权播放率、有效账号数。",
        "主图：账号消耗柱 + 累计消耗占比折线；加入 80% 参考线。",
        "说明卡：单月、深度脱敏、同比例缩放、ROI=1 非利润线。",
    ):
        add_bullet(document, item)

    add_heading(document, "6.4 第 2 页：品牌对标", 2)
    for item in (
        "散点：X=总消耗，Y=24H 混合 ROI，大小=账号数，图例=品牌；加入 ROI=1 参考线。",
        "水平条形图：品牌总消耗降序。",
        "矩阵：总消耗、混合 ROI、加权 CTR/CPC/播放率、重算播放成本。",
        "解释：品牌 B/E 进入排查，不直接停投；品牌 A 样本太小，不作稳定判断。",
    ):
        add_bullet(document, item)

    add_heading(document, "6.5 第 3 页：账号行动池", 2)
    for item in (
        "横条：五类行动标签账号数。",
        "明细：account_id、account_name、brand_name、platform_spend、mixed_roi_recalc、ctr_recalc、play_rate_recalc、play_cost_recalc、action_segment。",
        "按消耗降序；增加搜索切片器；颜色只是辅助，标签文字必须完整。",
        "页面写明：消耗门槛 50、有效账号 P75=735.80、ROI 参考线 1.0。",
    ):
        add_bullet(document, item)
    add_image(document, "images/dashboard_action_pool.png", caption="静态 GitHub 看板的行动池页；Power BI 应保持同一口径与信息层级。")

    add_heading(document, "7. React 看板、本地 QA 与 GitHub Pages", 1)
    add_heading(document, "7.1 本地启动", 2)
    add_code_block(
        document,
        """cd web
npm ci
npm test
npm run dev

# 生产构建
npm run build
npm run preview""",
        language="shell",
    )
    add_body(document, "打开终端给出的 http://localhost 地址。经营总览、品牌对标、账号行动池三个页签都应可见；筛选后 KPI、图表和表格联动。")
    add_image(document, "images/dashboard_overview.png", caption="经营总览：KPI、Pareto、洞察轨道与方法边界。")

    add_heading(document, "7.2 手工验收", 2)
    for item in (
        "无筛选：总消耗 207,057.29、ROI 1.03、有效账号 283。",
        "单品牌：KPI 与图表同时变化，品牌矩阵只保留所选品牌。",
        "单行动标签：明细行数与标签计数一致。",
        "方法说明：能看到播放成本冲突、单月与脱敏边界。",
        "移动端 390px：页面无全局横向滚动；明细表只在自己的容器内滚动。",
        "浏览器控制台无 JavaScript 错误，Network 中 dashboard.json 返回 200。",
    ):
        add_bullet(document, item)

    add_heading(document, "7.3 数据更新", 2)
    add_body(document, "网页只读取 web/public/data/dashboard.json。不要手工编辑 JSON；每次通过 Python 管道生成，才能同时更新 CSV、图、Notebook 与对账回执。")
    add_code_block(
        document,
        """python scripts/run_pipeline.py
python scripts/build_notebooks.py
python scripts/execute_notebooks.py
PYTHONPATH=src python -m unittest discover -s tests -v
cd web && npm test && npm run build && cd ..""",
        language="shell",
    )

    add_heading(document, "8. 结果解读、行动框架与面试表达", 1)
    add_heading(document, "8.1 五类行动池", 2)
    add_table(
        document,
        ["标签", "账号数", "建议动作", "必须验证"],
        [
            ["核心扩量候选", "62", "小幅提额、分批观察", "增量消耗是否仍保持回收"],
            ["高消耗重点优化", "9", "先查配置/素材/承接，必要时降额", "先排除数据与配置异常"],
            ["小步扩量观察", "137", "补量跨门槛，验证稳定性", "低消耗 ROI 波动"],
            ["低效清理候选", "75", "复核后降级或暂停", "不做机械批量清理"],
            ["数据不足/低量池", "268", "补样本、保留测试或归档", "不能直接判定好坏"],
        ],
        [1.55, 0.75, 2.2, 2.0],
    )

    add_heading(document, "8.2 三条业务结论", 2)
    for item in (
        "预算高度集中：Top 55 账号贡献 64.74% 消耗，因此运营资源应先管头部，再覆盖长尾。",
        "规模与效率要一起看：品牌 C 占 48.19% 消耗且 ROI 1.04；品牌 B/E 的 ROI 低于 1，但只能进入排查，不直接归因。",
        "单一前链路指标解释力弱：有效样本中 CTR/CPC/播放率与 ROI 的 Spearman 绝对值均低于 0.04；重算播放成本为 -0.128。",
    ):
        add_bullet(document, item)

    add_heading(document, "8.3 30 秒版本", 2)
    add_callout(
        document,
        "讲稿",
        "我基于实习场景的一份深度脱敏账号月度数据，主动完成了一个可复现的投放效能诊断项目。先锁定账号月粒度和加权口径，审计出源播放成本方向相反，再用 Python、MySQL、Excel 数据字典、React 看板和 Power BI 施工包形成同口径交付。最后以消耗门槛、有效账号 P75 与 ROI 参考线把 551 个账号分成五类行动池；项目只输出优先级，不声称因果或收益。",
        tone="teal",
    )

    add_heading(document, "8.4 面试追问底线", 2)
    add_table(
        document,
        ["追问", "回答核心"],
        [
            ["公司里实际做了什么？", "素材上传、B 站推送、广告系统品牌/产品基建与计划搭建；三次历史周复盘是主动额外完成。"],
            ["为什么不平均 ROI？", "账号消耗规模不同；组合 ROI 必须总变现 ÷ 总消耗。"],
            ["为什么门槛是 50？", "是透明初始门槛；283 个账号覆盖 99.43% 消耗，线上仍需敏感性验证。"],
            ["ROI>1 是否赚钱？", "否；缺少内容成本、分成、人力、税费、退款与长周期回收。"],
            ["为什么没有 PBIX？", "不伪造未亲手生成的文件；提供可审计模型、M、DAX、主题与构建清单。"],
            ["项目效果如何？", "可验证结果是完整行动池与可复现资产；没有证据支持业务提升。"],
        ],
        [2.0, 4.5],
        font_size=8.4,
    )

    add_heading(document, "9. Git/GitHub 上传与更新", 1)
    add_heading(document, "9.1 创建仓库并首次推送", 2)
    for step in (
        "GitHub 右上角 + → New repository，仓库名 manga-ad-account-analysis。",
        "选择 Public；不要勾选自动 README、.gitignore 或 License，因为项目已包含。",
        "在项目根目录初始化 Git、提交、切换 main，并关联远端。",
        "把 <你的用户名> 替换为真实用户名后推送。不要把尖括号一起保留。",
    ):
        add_step(document, step)
    add_code_block(
        document,
        """git init
git add .
git commit -m \"feat: publish manga ad account analysis portfolio\"
git branch -M main
git remote add origin https://github.com/<你的用户名>/manga-ad-account-analysis.git
git push -u origin main""",
        language="shell",
    )

    add_heading(document, "9.2 GitHub Pages 部署", 2)
    for step in (
        "仓库 Settings → Pages。",
        "Build and deployment → Source 选择 GitHub Actions。",
        "进入 Actions，打开 Deploy dashboard to GitHub Pages，等待 build 与 deploy 变绿。",
        "访问 https://<你的用户名>.github.io/manga-ad-account-analysis/。",
        "把仓库与 Pages 链接补充到你自己的简历版本；先打开验证，再投递。",
    ):
        add_step(document, step)
    add_callout(
        document,
        "隐私检查",
        "推送前执行 git status，并搜索 customer_id、customer_name、token、password、.env、原始公司内部名称与本机绝对路径。公开 JSON 只允许匿名账号、品牌和聚合指标。",
        tone="coral",
    )

    add_heading(document, "9.3 日常更新", 2)
    add_code_block(
        document,
        """git status
git add data/processed web/public/data images notebooks outputs
git commit -m \"data: refresh anonymized analysis outputs\"
git push""",
        language="shell",
    )

    add_heading(document, "10. 故障排查与最终验收", 1)
    add_heading(document, "10.1 常见故障", 2)
    add_table(
        document,
        ["症状", "原因", "处理"],
        [
            ["python 找不到", "未加 PATH / 终端未重开", "重新安装并勾选 PATH；重开终端"],
            ["Activate.ps1 被禁", "PowerShell 执行策略", "当前窗口 Set-ExecutionPolicy -Scope Process Bypass"],
            ["ModuleNotFoundError", "虚拟环境未激活或依赖未装", "激活 .venv；pip install -e \".[dev]\""],
            ["XLSX 工作表不存在", "源文件错误或工作表改名", "必须包含 脱敏数据；不要改名"],
            ["MySQL 导入乱码", "CSV 编码/分隔符不对", "使用 CSV UTF-8 与导入向导"],
            ["SQL P75 不一致", "导入行数/小数精度/门槛不同", "先核对 551 行、DECIMAL 与 spend>=50"],
            ["Power BI ROI 不一致", "AVERAGE 了账号 ROI", "改用 SUM(变现)/SUM(消耗)"],
            ["网页空白", "dashboard.json 未提交或 Pages 配置错", "Network 看 200；Pages Source 选 Actions"],
            ["移动端横向滚动", "明细表溢出到页面", "保留表格容器滚动，不让 body 溢出"],
        ],
        [1.6, 2.05, 2.85],
        font_size=7.9,
    )

    add_heading(document, "10.2 完成清单", 2)
    checks = (
        "[ ] Python 全部测试通过，pipeline_receipt.json 对账为 true。",
        "[ ] Notebook 从上到下执行，无错误单元。",
        "[ ] 数据字典 5 个工作表可读，无公式错误。",
        "[ ] 网页三页、品牌筛选、行动筛选与移动端通过。",
        "[ ] Power BI 无筛选基线与 Python 一致；三页均写明单月/脱敏边界。",
        "[ ] PDF 报告 6 页可读，无裁切或乱码。",
        "[ ] README、报告、看板、简历使用同一套数字与动作标签。",
        "[ ] 仓库没有客户字段、Token、口令、.env、原始未脱敏文件或本机绝对路径。",
        "[ ] 简历明确写“个人分析项目”，实习职责与主动复盘分开。",
        "[ ] 能在 3 分钟内解释口径、发现、动作、边界与下一步。",
    )
    for item in checks:
        add_bullet(document, item)

    add_heading(document, "附录 A. 完整代码附录", 1)
    add_body(document, "以下代码在生成本手册时直接读取自最终仓库，没有另写一个“示例版本”。复制时请以代码标题中的相对路径为准。前端样式与全部自动化测试仍保留在压缩包中；日常复现无需手工抄写。")

    code_files = [
        "pyproject.toml",
        "requirements.txt",
        ".gitignore",
        "src/manga_ad_analysis/__init__.py",
        "src/manga_ad_analysis/io.py",
        "src/manga_ad_analysis/quality.py",
        "src/manga_ad_analysis/metrics.py",
        "src/manga_ad_analysis/segmentation.py",
        "src/manga_ad_analysis/pipeline.py",
        "src/manga_ad_analysis/visuals.py",
        "scripts/run_pipeline.py",
        "scripts/build_notebooks.py",
        "scripts/execute_notebooks.py",
        "sql/01_create_table.sql",
        "sql/02_data_quality.sql",
        "sql/03_brand_summary.sql",
        "sql/04_account_pareto.sql",
        "sql/05_account_segmentation.sql",
        "powerbi/power_query.m",
        "powerbi/measures.dax",
        "powerbi/theme.json",
        "web/package.json",
        "web/vite.config.js",
        "web/src/main.jsx",
        "web/src/App.jsx",
        "web/src/lib/data.js",
        "web/src/pages/OverviewPage.jsx",
        "web/src/pages/BrandPage.jsx",
        "web/src/pages/ActionPage.jsx",
        ".github/workflows/deploy-pages.yml",
    ]
    language_by_suffix = {
        ".py": "python",
        ".sql": "sql",
        ".dax": "dax",
        ".m": "power-query-m",
        ".json": "json",
        ".jsx": "javascript",
        ".js": "javascript",
        ".yml": "yaml",
        ".yaml": "yaml",
        ".toml": "toml",
    }
    for relative_path in code_files:
        path = PROJECT_ROOT / relative_path
        if not path.exists():
            continue
        add_heading(document, relative_path, 2)
        add_code_block(
            document,
            path.read_text(encoding="utf-8"),
            language=language_by_suffix.get(path.suffix.lower(), "text"),
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def configure_resume(document: Document) -> None:
    section = document.sections[0]
    # Named override: A4 + compact margins are appropriate for a Chinese one-page resume.
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.46)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(8.6)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(1.8)
    normal.paragraph_format.line_spacing = 1.05

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(8.25)
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.first_line_indent = Inches(-0.13)
        style.paragraph_format.space_after = Pt(1.0)
        style.paragraph_format.line_spacing = 1.02


def resume_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3.5)
    paragraph.paragraph_format.space_after = Pt(2.5)
    add_paragraph_border(paragraph, color=BLUE, size="8")
    set_run_font(paragraph.add_run(title), size=10, color=NAVY, bold=True)


def resume_entry(document: Document, left: str, right: str = "", *, subtitle: str = "") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1.0)
    paragraph.paragraph_format.space_after = Pt(1.2)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.05), alignment=WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(paragraph.add_run(left), size=8.9, color=INK, bold=True)
    if right:
        set_run_font(paragraph.add_run("\t" + right), size=8.2, color=MUTED, bold=True)
    if subtitle:
        set_run_font(paragraph.add_run("  " + subtitle), size=8.2, color=MUTED)


def resume_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0.8)
    paragraph.paragraph_format.line_spacing = 1.02
    set_run_font(paragraph.add_run(text), size=8.15)


def build_resume(output_path: Path) -> None:
    document = Document()
    configure_resume(document)
    properties = document.core_properties
    properties.title = "蔡健明｜字节数据岗位校招简历"
    properties.subject = "数据分析/经营分析/数据策略/数据运营"
    properties.author = "蔡健明"
    properties.last_modified_by = "蔡健明"

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    set_run_font(name.add_run("蔡健明"), size=25, color=NAVY, bold=True)

    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target.paragraph_format.space_after = Pt(1.5)
    set_run_font(target.add_run("2027 届｜数据分析 / 数据策略 / 经营分析"), size=9.6, color=CORAL, bold=True)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(2.5)
    set_run_font(
        contact.add_run("19012716053  ·  1764798435@qq.com  ·  杭州"),
        size=8.4,
        color=CORAL,
        bold=True,
    )
    add_paragraph_border(contact, color=CORAL, size="10")

    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2.4)
    profile.paragraph_format.line_spacing = 1.03
    set_run_font(
        profile.add_run("经济统计学本科在读，具备 SQL、Python、Excel 与 Power BI 的完整分析链路实践；能从数据质量、加权口径、分层诊断推进到可复现看板与业务表达。"),
        size=8.35,
        color=INK,
    )

    resume_section(document, "教育背景")
    resume_entry(document, "浙江财经大学｜经济统计学 本科", "2023.09–2027.06")
    detail = document.add_paragraph()
    detail.paragraph_format.space_after = Pt(1.2)
    set_run_font(detail.add_run("GPA 3.26/4.00｜CET-6｜核心课程：Python、R、SQL 数据库、多元统计、时间序列、计量经济学、数据挖掘"), size=8.15, color=MUTED)

    resume_section(document, "核心项目")
    resume_entry(document, "漫剧投放账号效能诊断与预算分层｜个人分析项目", "2026.07–2026.08")
    resume_bullet(document, "基于 2026 年 7 月 551 个匿名账号、8 个品牌的账号月度数据，定义加权 CTR/CPC/播放率/ROI，完成缺失、重复、零分母与字段口径审计。")
    resume_bullet(document, "发现源“播放成本”实际匹配播放量÷消耗，统一重算为消耗÷播放量；以消耗门槛 50、有效账号消耗 P75=735.80、ROI 参考线 1.0 建立五类互斥行动池。")
    resume_bullet(document, "使用 Python/pandas、MySQL 8 CTE/窗口函数、Excel 数据字典与执行 Notebook 形成可复现管道；Top 55 账号贡献 64.74% 消耗，283 个有效账号覆盖 99.43% 消耗。")
    resume_bullet(document, "搭建 React/Vite 静态交互看板与 GitHub Pages 工作流，并提供 Power BI 星型模型、Power Query、DAX、主题和三页验收清单；明确单月数据不支持趋势、因果或策略收益结论。")

    resume_section(document, "实习经历")
    resume_entry(document, "掌玩｜运营实习生", "2026.05–至今")
    resume_bullet(document, "负责视频素材上传与 B 站内容推送，参与广告系统品牌/产品基础配置及投放计划搭建，按要求完成执行记录与异常反馈。")
    resume_bullet(document, "在日常工作外主动使用 Excel、基础 R/Python 与自建前端完成三次历史周数据复盘并向带教展示；相关分析为自主复盘，不作为公司正式岗位成果表述。")

    resume_entry(document, "宁德新能源科技有限公司｜PMC 数据实习生", "2025.07–2025.08")
    resume_bullet(document, "协助整理、核对并导入生产数据，参与排程系统测试；通过 MES、飞书与 Excel 查询统计生产进度，记录异常并及时反馈。")
    resume_bullet(document, "制作基础可视化报表，支持团队查看生产节奏与问题项；在数据口径不一致时回到原始记录核验。")

    resume_section(document, "其他项目与校园经历")
    resume_bullet(document, "浙江省城乡居民收入分析：完成数据清洗、城乡对比与可视化；个人数据分析静态网站：独立完成需求梳理、数据处理、页面开发、测试与部署。")
    resume_bullet(document, "数据科学学院文艺团副团长（2024.06–2025.06），协调组织 10+ 场活动；街舞队成员（2023.09–至今），获浙江省大学生艺术节团体一等奖等奖项。")

    resume_section(document, "技能")
    skills = document.add_paragraph()
    skills.paragraph_format.space_after = Pt(0)
    skills.paragraph_format.line_spacing = 1.03
    set_run_font(skills.add_run("数据工具："), size=8.2, color=NAVY, bold=True)
    set_run_font(skills.add_run("MySQL/SQL（CTE、窗口函数、聚合、CASE）、Python（pandas/NumPy/Matplotlib）、Excel、R；"), size=8.2)
    set_run_font(skills.add_run("BI 与工程："), size=8.2, color=NAVY, bold=True)
    set_run_font(skills.add_run("Power BI（DAX、Power Query、星型模型）、React/Vite、Git/GitHub Pages；"), size=8.2)
    set_run_font(skills.add_run("分析方法："), size=8.2, color=NAVY, bold=True)
    set_run_font(skills.add_run("数据质量、加权指标、Pareto、分层诊断、可复现分析与业务边界表达。"), size=8.2)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_resume_html(output_path: Path) -> None:
    regular_font = (FONT_DIR / "NotoSansCJKsc-Regular.otf").as_uri()
    bold_font = (FONT_DIR / "NotoSansCJKsc-Bold.otf").as_uri()
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>蔡健明｜字节数据岗位校招简历</title>
  <style>
    @font-face {{ font-family: NotoCJK; src: url('{regular_font}') format('opentype'); font-weight: 400; }}
    @font-face {{ font-family: NotoCJK; src: url('{bold_font}') format('opentype'); font-weight: 700; }}
    @page {{ size: A4; margin: 0; }}
    * {{ box-sizing: border-box; }}
    html, body {{ margin: 0; padding: 0; background: #fff; color: #17213a; font-family: NotoCJK, sans-serif; }}
    .page {{ width: 210mm; height: 297mm; padding: 10.5mm 14.5mm 9.5mm; overflow: hidden; }}
    h1 {{ margin: 0; text-align: center; font-size: 25pt; line-height: 1.05; color: #1b2a4a; }}
    .target {{ text-align: center; margin: 1.2mm 0 0; color: #e86a4a; font-size: 9.6pt; font-weight: 700; }}
    .contact {{ text-align: center; margin: 1mm 0 2.2mm; padding-bottom: 2.2mm; border-bottom: 1.8px solid #e86a4a; color: #e86a4a; font-size: 8.2pt; font-weight: 700; }}
    .profile {{ margin: 0 0 1.6mm; font-size: 8.2pt; line-height: 1.35; }}
    h2 {{ margin: 2.1mm 0 1.1mm; padding-bottom: .8mm; border-bottom: 1.2px solid #3867d6; font-size: 10pt; line-height: 1.1; color: #1b2a4a; }}
    .entry {{ display: grid; grid-template-columns: 1fr auto; gap: 4mm; align-items: baseline; margin: .8mm 0 .6mm; }}
    .entry strong {{ font-size: 8.7pt; }}
    .entry time {{ color: #667085; font-size: 8pt; font-weight: 700; white-space: nowrap; }}
    .detail {{ margin: 0; color: #667085; font-size: 8pt; line-height: 1.3; }}
    ul {{ margin: .3mm 0 0; padding-left: 4.2mm; }}
    li {{ margin: 0 0 .7mm; padding-left: .4mm; font-size: 7.9pt; line-height: 1.32; }}
    .skills {{ margin: 0; font-size: 8pt; line-height: 1.36; }}
    .skills b {{ color: #1b2a4a; }}
    .note {{ margin: 1.5mm 0 0; text-align: right; color: #e86a4a; font-size: 7.2pt; }}
  </style>
</head>
<body><main class="page">
  <h1>蔡健明</h1>
  <p class="target">2027 届｜数据分析 / 数据策略 / 经营分析</p>
  <p class="contact">19012716053 · 1764798435@qq.com · 杭州</p>
  <p class="profile">经济统计学本科在读，具备 SQL、Python、Excel 与 Power BI 的完整分析链路实践；能从数据质量、加权口径、分层诊断推进到可复现看板与业务表达。</p>

  <h2>教育背景</h2>
  <div class="entry"><strong>浙江财经大学｜经济统计学 本科</strong><time>2023.09–2027.06</time></div>
  <p class="detail">GPA 3.26/4.00｜CET-6｜核心课程：Python、R、SQL 数据库、多元统计、时间序列、计量经济学、数据挖掘</p>

  <h2>核心项目</h2>
  <div class="entry"><strong>漫剧投放账号效能诊断与预算分层｜个人分析项目</strong><time>2026.07–2026.08</time></div>
  <ul>
    <li>基于 2026 年 7 月 551 个匿名账号、8 个品牌的账号月度数据，定义加权 CTR/CPC/播放率/ROI，完成缺失、重复、零分母与字段口径审计。</li>
    <li>发现源“播放成本”实际匹配播放量÷消耗，统一重算为消耗÷播放量；以消耗门槛 50、有效账号消耗 P75=735.80、ROI 参考线 1.0 建立五类互斥行动池。</li>
    <li>使用 Python/pandas、MySQL 8 CTE/窗口函数、Excel 数据字典与执行 Notebook 形成可复现管道；Top 55 账号贡献 64.74% 消耗，283 个有效账号覆盖 99.43% 消耗。</li>
    <li>搭建 React/Vite 静态交互看板与 GitHub Pages 工作流，并提供 Power BI 星型模型、Power Query、DAX、主题和三页验收清单；明确单月数据不支持趋势、因果或策略收益结论。</li>
  </ul>

  <h2>实习经历</h2>
  <div class="entry"><strong>掌玩｜运营实习生</strong><time>2026.05–至今</time></div>
  <ul>
    <li>负责视频素材上传与 B 站内容推送，参与广告系统品牌/产品基础配置及投放计划搭建，按要求完成执行记录与异常反馈。</li>
    <li>在日常工作外主动使用 Excel、基础 R/Python 与自建前端完成三次历史周数据复盘并向带教展示；相关分析为自主复盘，不作为公司正式岗位成果表述。</li>
  </ul>
  <div class="entry"><strong>宁德新能源科技有限公司｜PMC 数据实习生</strong><time>2025.07–2025.08</time></div>
  <ul>
    <li>协助整理、核对并导入生产数据，参与排程系统测试；通过 MES、飞书与 Excel 查询统计生产进度，记录异常并及时反馈。</li>
    <li>制作基础可视化报表，支持团队查看生产节奏与问题项；在数据口径不一致时回到原始记录核验。</li>
  </ul>

  <h2>其他项目与校园经历</h2>
  <ul>
    <li>浙江省城乡居民收入分析：完成数据清洗、城乡对比与可视化；个人数据分析静态网站：独立完成需求梳理、数据处理、页面开发、测试与部署。</li>
    <li>数据科学学院文艺团副团长（2024.06–2025.06），协调组织 10+ 场活动；街舞队成员（2023.09–至今），获浙江省大学生艺术节团体一等奖等奖项。</li>
  </ul>

  <h2>技能</h2>
  <p class="skills"><b>数据工具：</b>MySQL/SQL（CTE、窗口函数、聚合、CASE）、Python（pandas/NumPy/Matplotlib）、Excel、R；<b>BI 与工程：</b>Power BI（DAX、Power Query、星型模型）、React/Vite、Git/GitHub Pages；<b>分析方法：</b>数据质量、加权指标、Pareto、分层诊断、可复现分析与业务边界表达。</p>
</main></body></html>"""
    output_path.write_text(html, encoding="utf-8")


def export_resume_pdf(resume_docx: Path, output_pdf: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="resume-pdf-") as directory:
        temp_root = Path(directory)
        html_path = temp_root / "resume.html"
        build_resume_html(html_path)
        node = os.environ.get("CODEX_PRIMARY_RUNTIME_NODE") or shutil.which("node")
        if not node:
            raise RuntimeError("Node.js is required to export the resume PDF")
        environment = os.environ.copy()
        home = temp_root / "home"
        cache = temp_root / "cache"
        home.mkdir()
        cache.mkdir()
        environment["HOME"] = str(home)
        environment["TMPDIR"] = str(temp_root)
        environment["XDG_CACHE_HOME"] = str(cache)
        completed = subprocess.run(
            [
                node,
                str(PROJECT_ROOT / "tools" / "print_to_pdf.mjs"),
                str(html_path),
                str(output_pdf),
            ],
            capture_output=True,
            text=True,
            env=environment,
            check=False,
        )
        if completed.returncode != 0 or not output_pdf.exists():
            raise RuntimeError(f"Resume PDF export failed: {completed.stderr or completed.stdout}")


def build_all(output_dir: Path) -> dict[str, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    manual = output_dir / MANUAL_NAME
    resume_docx = output_dir / RESUME_DOCX_NAME
    resume_pdf = output_dir / RESUME_PDF_NAME
    build_manual(manual)
    build_resume(resume_docx)
    export_resume_pdf(resume_docx, resume_pdf)
    return {
        "manual": str(manual),
        "resume_docx": str(resume_docx),
        "resume_pdf": str(resume_pdf),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    result = build_all(args.output_dir.resolve())
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
